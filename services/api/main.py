"""
FastAPI application — the HTTP/WebSocket gateway for the research pipeline.

Exposes endpoints for authentication (email + Google OAuth), launching and
polling research jobs, live agent-status streaming over WebSocket, quick RAG
Q&A, document upload and Q&A, report refinement, audit-log access, exporting
reports to Notion / Google Docs, and notifications. Business logic lives in the
agent modules and the shared package; this module wires HTTP I/O to them.
"""
from __future__ import annotations
import asyncio
import json
import logging
import uuid
from contextlib import asynccontextmanager
from typing import Optional

from fastapi import FastAPI, WebSocket, WebSocketDisconnect, HTTPException, UploadFile, File, Form, Depends, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse, StreamingResponse
from pydantic import BaseModel, EmailStr

from shared.database import init_db, SessionLocal, ResearchJob, AgentTask, User, Document, ChatHistory, QAInteraction, ClaimVerification, ReasoningTrace, NotificationLog, OAuthConnection, AuditLog
from shared.message_bus import bus
from shared.logging_config import configure_logging
from shared.auth import (
    hash_password, verify_password, create_access_token,
    create_verification_token, decode_token, send_verification_email,
)
from agents.orchestrator import run_research

configure_logging()
logger = logging.getLogger(__name__)


@asynccontextmanager
async def lifespan(app: FastAPI):
    configure_logging()
    init_db()
    logger.info("API startup complete — logging and database initialized")
    yield
    logger.info("API shutting down")

app = FastAPI(title="Research Pipeline API", lifespan=lifespan)

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_methods=["*"],
    allow_headers=["*"],
)


# ── Auth helpers ──────────────────────────────────────────────────────────────

def get_current_user(request: Request) -> Optional[dict]:
    """Extract user from Authorization header or ?token= query param. Returns None if not authenticated."""
    auth = request.headers.get("Authorization", "")
    if auth.startswith("Bearer "):
        token = auth[7:]
    else:
        # Fallback: check query param (for browser redirects like OAuth connect)
        token = request.query_params.get("token", "")
    if not token:
        return None
    payload = decode_token(token)
    if not payload or "sub" not in payload:
        return None
    return {"user_id": payload["sub"], "email": payload.get("email", "")}


def require_auth(request: Request) -> dict:
    """Dependency that requires authentication."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    return user


# ── Auth endpoints ────────────────────────────────────────────────────────────

class SignupRequest(BaseModel):
    email: str
    password: str
    name: str = ""


class LoginRequest(BaseModel):
    email: str
    password: str


@app.post("/auth/signup")
async def signup(req: SignupRequest):
    """Register a new user with email + password. Sends verification email."""
    if not req.email or not req.password:
        raise HTTPException(status_code=400, detail="Email and password required")
    if len(req.password) < 6:
        raise HTTPException(status_code=400, detail="Password must be at least 6 characters")

    db = SessionLocal()
    existing = db.query(User).filter(User.email == req.email).first()
    if existing:
        db.close()
        raise HTTPException(status_code=409, detail="Email already registered")

    user = User(
        email=req.email,
        password_hash=hash_password(req.password),
        name=req.name,
        provider="email",
        is_verified=False,
    )
    db.add(user)
    db.commit()
    user_id = user.id
    db.close()

    # Send verification email
    token = create_verification_token(req.email)
    await send_verification_email(req.email, token)

    return {"message": "Account created. Check your email to verify.", "user_id": user_id}


@app.post("/auth/login")
async def login(req: LoginRequest):
    """Login with email + password. Returns JWT token."""
    db = SessionLocal()
    user = db.query(User).filter(User.email == req.email).first()
    db.close()

    if not user or not user.password_hash:
        raise HTTPException(status_code=401, detail="Invalid email or password")
    if not verify_password(req.password, user.password_hash):
        raise HTTPException(status_code=401, detail="Invalid email or password")
    if not user.is_verified:
        raise HTTPException(status_code=403, detail="Email not verified. Check your inbox.")

    token = create_access_token(user.id, user.email)
    return {"token": token, "user": {"id": user.id, "email": user.email, "name": user.name}}


@app.get("/auth/verify")
async def verify_email(token: str):
    """Verify email via link. Redirects to UI on success."""
    payload = decode_token(token)
    if not payload or payload.get("purpose") != "verify":
        raise HTTPException(status_code=400, detail="Invalid or expired verification link")

    email = payload.get("email")
    db = SessionLocal()
    user = db.query(User).filter(User.email == email).first()
    if user:
        user.is_verified = True
        db.commit()
    db.close()

    import os
    base_url = os.getenv("APP_BASE_URL", "http://localhost:3000")
    return RedirectResponse(url=f"{base_url}/?verified=true")


@app.get("/auth/me")
async def get_me(user: dict = Depends(require_auth)):
    """Get current user info."""
    db = SessionLocal()
    db_user = db.query(User).filter(User.id == user["user_id"]).first()
    db.close()
    if not db_user:
        raise HTTPException(status_code=404, detail="User not found")
    return {"id": db_user.id, "email": db_user.email, "name": db_user.name}


# ── Google OAuth ──────────────────────────────────────────────────────────────

@app.get("/auth/google")
async def google_login():
    """Redirect to Google OAuth consent screen."""
    import os
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    if not client_id:
        raise HTTPException(status_code=500, detail="Google OAuth not configured")

    base_url = os.getenv("APP_BASE_URL", "http://localhost:3000")
    redirect_uri = f"{base_url}/api/auth/google/callback"
    scope = "openid email profile"
    url = (
        f"https://accounts.google.com/o/oauth2/v2/auth"
        f"?client_id={client_id}"
        f"&redirect_uri={redirect_uri}"
        f"&response_type=code"
        f"&scope={scope}"
        f"&access_type=offline"
    )
    return RedirectResponse(url=url)


@app.get("/auth/google/callback")
async def google_callback(code: str):
    """Handle Google OAuth callback."""
    import os
    import httpx

    client_id = os.getenv("GOOGLE_CLIENT_ID")
    client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
    base_url = os.getenv("APP_BASE_URL", "http://localhost:3000")
    redirect_uri = f"{base_url}/api/auth/google/callback"

    # Exchange code for tokens
    async with httpx.AsyncClient() as http:
        token_res = await http.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": client_id,
                "client_secret": client_secret,
                "redirect_uri": redirect_uri,
                "grant_type": "authorization_code",
            },
        )
        if token_res.status_code != 200:
            raise HTTPException(status_code=400, detail="Google auth failed")
        tokens = token_res.json()

        # Get user info
        userinfo_res = await http.get(
            "https://www.googleapis.com/oauth2/v2/userinfo",
            headers={"Authorization": f"Bearer {tokens['access_token']}"},
        )
        userinfo = userinfo_res.json()

    email = userinfo.get("email")
    name = userinfo.get("name", "")

    if not email:
        raise HTTPException(status_code=400, detail="Could not get email from Google")

    # Find or create user
    db = SessionLocal()
    user = db.query(User).filter(User.email == email).first()
    if not user:
        user = User(email=email, name=name, provider="google", is_verified=True)
        db.add(user)
        db.commit()
    elif not user.is_verified:
        user.is_verified = True
        db.commit()
    user_id = user.id
    db.close()

    # Create JWT and redirect to UI with token
    token = create_access_token(user_id, email)
    return RedirectResponse(url=f"{base_url}/?token={token}")


# ── REST endpoints ────────────────────────────────────────────────────────────

class ResearchRequest(BaseModel):
    query: str
    language: str = "English"


@app.post("/research")
async def start_research(req: ResearchRequest, request: Request):
    """Kick off a research job. Returns immediately with a job_id."""
    if not req.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    user = get_current_user(request)
    user_id = user["user_id"] if user else None

    import uuid
    job_id = str(uuid.uuid4())

    # Save to chat history
    if user_id:
        title = await _generate_chat_title(req.query)
        _save_chat_history(user_id, "Research", title, job_id)

    # Run pipeline in background so the HTTP response returns immediately
    async def _run():
        try:
            await run_research(req.query, job_id=job_id, user_id=user_id, language=req.language)
        except Exception as e:
            logger.error("Background pipeline error for job %s: %s", job_id, e, exc_info=True)

    asyncio.create_task(_run())

    return {"job_id": job_id, "status": "running", "query": req.query}


@app.get("/research/{job_id}")
async def get_job(job_id: str):
    """Poll job status and retrieve the report when done."""
    db = SessionLocal()
    job = db.query(ResearchJob).filter(ResearchJob.id == job_id).first()
    db.close()

    if not job:
        raise HTTPException(status_code=404, detail="Job not found")

    return {
        "job_id": job.id,
        "query": job.query,
        "status": job.status,
        "report": job.report,
        "created_at": job.created_at,
        "finished_at": job.finished_at,
    }


@app.get("/research/{job_id}/tasks")
async def get_tasks(job_id: str):
    """Return all agent task records for a job — useful for the audit trail."""
    db = SessionLocal()
    tasks = db.query(AgentTask).filter(AgentTask.job_id == job_id).all()
    db.close()

    return [
        {
            "agent": t.agent_name,
            "status": t.status,
            "duration_ms": round(t.duration_ms or 0),
            "started_at": t.started_at,
            "finished_at": t.finished_at,
        }
        for t in tasks
    ]


@app.get("/research/{job_id}/claims")
async def get_claims(job_id: str):
    """Return all fact-check claim verifications for a job."""
    db = SessionLocal()
    claims = db.query(ClaimVerification).filter(ClaimVerification.job_id == job_id).all()
    db.close()

    return [
        {
            "id": c.id,
            "claim": c.claim,
            "status": c.status,
            "confidence": c.confidence,
            "supported_by": json.loads(c.supported_by) if c.supported_by else [],
            "contradicted_by": json.loads(c.contradicted_by) if c.contradicted_by else [],
            "sentence_match": c.sentence_match,
        }
        for c in claims
    ]


@app.get("/research/{job_id}/reasoning")
async def get_reasoning(job_id: str):
    """Return all agent reasoning traces for a job."""
    db = SessionLocal()
    traces = db.query(ReasoningTrace).filter(ReasoningTrace.job_id == job_id).order_by(ReasoningTrace.created_at).all()
    db.close()

    return [
        {
            "id": t.id,
            "agent": t.agent,
            "step": t.step,
            "reasoning": t.reasoning,
            "decision": t.decision,
            "metadata": json.loads(t.trace_metadata) if t.trace_metadata else {},
            "created_at": t.created_at.isoformat() if t.created_at else None,
        }
        for t in traces
    ]


@app.get("/research/{job_id}/audit")
async def get_audit_log(job_id: str, level: Optional[str] = None):
    """Return the audit log for a job — what ran, what failed, and why."""
    db = SessionLocal()
    try:
        q = db.query(AuditLog).filter(AuditLog.job_id == job_id)
        if level:
            q = q.filter(AuditLog.level == level)
        entries = q.order_by(AuditLog.created_at).all()
    finally:
        db.close()

    return [
        {
            "id": a.id,
            "level": a.level,
            "event": a.event,
            "agent": a.agent,
            "message": a.message,
            "detail": a.detail,
            "user_id": a.user_id,
            "created_at": a.created_at.isoformat() if a.created_at else None,
        }
        for a in entries
    ]


@app.get("/audit")
async def list_audit_log(
    user: dict = Depends(require_auth),
    user_id: Optional[str] = None,
    level: Optional[str] = None,
    event: Optional[str] = None,
    limit: int = 200,
):
    """
    Support/diagnostics endpoint: browse recent audit entries across jobs.
    Filter by user_id, level (info|warning|error), or event.
    """
    limit = max(1, min(limit, 1000))
    db = SessionLocal()
    try:
        q = db.query(AuditLog)
        if user_id:
            q = q.filter(AuditLog.user_id == user_id)
        if level:
            q = q.filter(AuditLog.level == level)
        if event:
            q = q.filter(AuditLog.event == event)
        entries = q.order_by(AuditLog.created_at.desc()).limit(limit).all()
    finally:
        db.close()

    return [
        {
            "id": a.id,
            "job_id": a.job_id,
            "user_id": a.user_id,
            "level": a.level,
            "event": a.event,
            "agent": a.agent,
            "message": a.message,
            "detail": a.detail,
            "created_at": a.created_at.isoformat() if a.created_at else None,
        }
        for a in entries
    ]


@app.get("/research/{job_id}/download")
async def download_report(job_id: str, format: str = "pdf"):
    """Download research report as PDF or DOCX."""
    db = SessionLocal()
    job = db.query(ResearchJob).filter(ResearchJob.id == job_id).first()
    db.close()

    if not job or not job.report:
        raise HTTPException(status_code=404, detail="Report not found")

    import io
    import markdown

    html_content = markdown.markdown(job.report, extensions=["tables", "fenced_code"])

    if format == "docx":
        from docx import Document as DocxDocument
        from docx.shared import Pt
        from bs4 import BeautifulSoup

        doc = DocxDocument()
        soup = BeautifulSoup(html_content, "html.parser")

        for element in soup.children:
            if element.name in ("h1", "h2", "h3"):
                level = int(element.name[1])
                doc.add_heading(element.get_text(), level=level)
            elif element.name == "p":
                doc.add_paragraph(element.get_text())
            elif element.name in ("ul", "ol"):
                for li in element.find_all("li"):
                    doc.add_paragraph(li.get_text(), style="List Bullet")
            elif element.name == "table":
                rows = element.find_all("tr")
                if rows:
                    cols = len(rows[0].find_all(["th", "td"]))
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = "Table Grid"
                    for i, row in enumerate(rows):
                        cells = row.find_all(["th", "td"])
                        for j, cell in enumerate(cells):
                            if j < cols:
                                table.rows[i].cells[j].text = cell.get_text()
            elif hasattr(element, "get_text") and element.get_text(strip=True):
                doc.add_paragraph(element.get_text())

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        filename = f"research_report_{job_id[:8]}.docx"
        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )

    else:  # PDF
        from weasyprint import HTML

        styled_html = f"""
        <html><head><style>
            body {{ font-family: 'Helvetica', 'Arial', sans-serif; font-size: 11px; line-height: 1.6; padding: 20px; }}
            h1 {{ font-size: 18px; }} h2 {{ font-size: 15px; }} h3 {{ font-size: 13px; }}
            table {{ border-collapse: collapse; width: 100%; margin: 10px 0; }}
            th, td {{ border: 1px solid #ccc; padding: 6px 8px; text-align: left; font-size: 10px; }}
            th {{ background: #f0f0f0; }}
        </style></head><body>{html_content}</body></html>
        """
        buffer = io.BytesIO()
        HTML(string=styled_html).write_pdf(buffer)
        buffer.seek(0)
        filename = f"research_report_{job_id[:8]}.pdf"
        return StreamingResponse(
            buffer,
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )


@app.get("/jobs")
async def list_jobs(request: Request, limit: int = 20):
    """List recent research jobs for the current user."""
    user = get_current_user(request)
    db = SessionLocal()
    query = db.query(ResearchJob).order_by(ResearchJob.created_at.desc())
    if user:
        query = query.filter(ResearchJob.user_id == user["user_id"])
    jobs = query.limit(limit).all()
    db.close()
    return [
        {"job_id": j.id, "query": j.query, "status": j.status, "created_at": j.created_at}
        for j in jobs
    ]


# ── Report Refinement (SSE streaming) ────────────────────────────────────────

class RefineRequest(BaseModel):
    instruction: str
    conversation_history: list = []  # [{role, content}]


@app.post("/research/{job_id}/refine")
async def refine_report(job_id: str, req: RefineRequest, request: Request):
    """Refine an existing report based on user instructions. Streams via SSE."""
    if not req.instruction.strip():
        raise HTTPException(status_code=400, detail="Instruction cannot be empty")

    db = SessionLocal()
    job = db.query(ResearchJob).filter(ResearchJob.id == job_id).first()
    if not job or not job.report:
        db.close()
        raise HTTPException(status_code=404, detail="Report not found")
    current_report = job.report
    db.close()

    from langchain_openai import AzureChatOpenAI
    from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
    import os

    llm = AzureChatOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-12-01-preview"),
        max_tokens=4000,
        streaming=True,
    )

    messages = [
        SystemMessage(content=(
            "You are an expert research report editor. The user has a research report and wants to modify it. "
            "Analyze their instruction carefully and produce the COMPLETE updated report in markdown format. "
            "If they ask to add a section, add it in the appropriate place. "
            "If they ask to remove something, remove it. "
            "If they ask to modify/rewrite something, do so while keeping the rest intact. "
            "If they have a general question about the report, answer it concisely then provide the full updated report. "
            "Always output the full updated report unless the user is just asking a question (not requesting changes). "
            "For questions that don't require report changes, just answer the question directly."
        )),
    ]

    # Add conversation history
    for msg in req.conversation_history:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        elif msg["role"] == "assistant":
            messages.append(AIMessage(content=msg["content"]))

    messages.append(HumanMessage(content=(
        f"Here is the current report:\n\n{current_report}\n\n"
        f"User instruction: {req.instruction}"
    )))

    async def event_generator():
        answer_text = ""
        try:
            yield f"data: {json.dumps({'type': 'start'})}\n\n"

            async for chunk in llm.astream(messages):
                token = chunk.content
                if token:
                    answer_text += token
                    yield f"data: {json.dumps({'type': 'token', 'token': token})}\n\n"

            # Determine if the response is a full updated report or just an answer
            is_report_update = (
                answer_text.strip().startswith("#") or
                "---" in answer_text[:100] or
                len(answer_text) > len(current_report) * 0.5
            )

            if is_report_update:
                # Save updated report to DB
                db2 = SessionLocal()
                try:
                    j = db2.query(ResearchJob).filter(ResearchJob.id == job_id).first()
                    if j:
                        j.report = answer_text
                        db2.commit()
                except Exception:
                    logger.error("Failed to save refined report for job %s", job_id, exc_info=True)
                    db2.rollback()
                finally:
                    db2.close()

            # Save the refinement interaction
            user = get_current_user(request)
            if user and answer_text:
                db3 = SessionLocal()
                try:
                    qa = QAInteraction(
                        user_id=user["user_id"],
                        chat_session_id=job_id,
                        document_id=None,
                        question=req.instruction,
                        answer=answer_text if not is_report_update else "[Report updated]",
                        sources=json.dumps({"type": "refinement", "job_id": job_id, "is_report_update": is_report_update}),
                    )
                    db3.add(qa)
                    db3.commit()
                except Exception:
                    logger.error("Failed to save refinement interaction for job %s", job_id, exc_info=True)
                    db3.rollback()
                finally:
                    db3.close()

            yield f"data: {json.dumps({'type': 'done', 'is_report_update': is_report_update})}\n\n"
        except Exception as e:
            logger.error("Refinement stream failed for job %s: %s", job_id, e, exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"

    return StreamingResponse(event_generator(), media_type="text/event-stream")


# ── Chat History ──────────────────────────────────────────────────────────────

import datetime as _dt


async def _generate_chat_title(question: str) -> str:
    """Use LLM to generate a short chat title from the user's question."""
    try:
        from langchain_openai import AzureChatOpenAI
        from langchain_core.messages import SystemMessage, HumanMessage
        import os

        llm = AzureChatOpenAI(
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-12-01-preview"),
            max_tokens=20,
        )
        resp = await llm.ainvoke([
            SystemMessage(content=(
                "Generate a very short title (3-6 words max) that summarizes the user's question or intent. "
                "No quotes, no punctuation at the end. Just the title."
            )),
            HumanMessage(content=question),
        ])
        title = resp.content.strip().strip('"').strip("'")
        return title[:60] if title else question[:60]
    except Exception:
        # Fallback: truncate the question itself
        logger.warning("Chat title generation failed; using truncated question", exc_info=True)
        return question[:60] if len(question) > 60 else question


def _save_chat_history(user_id: str, chat_type: str, title: str, ref_id: str):
    """Persist a chat history entry."""
    db = SessionLocal()
    try:
        entry = ChatHistory(
            user_id=user_id,
            type=chat_type,
            title=title,
            ref_id=ref_id,
            created_at=_dt.datetime.now(_dt.timezone(_dt.timedelta(hours=5, minutes=30))).replace(tzinfo=None),
        )
        db.add(entry)
        db.commit()
    except Exception:
        logger.error("Failed to persist chat history for user %s", user_id, exc_info=True)
        db.rollback()
    finally:
        db.close()


async def _generate_followups(llm, question: str, answer: str) -> list[str]:
    """Generate 3 follow-up questions based on the Q&A exchange."""
    try:
        from langchain_core.messages import SystemMessage, HumanMessage
        resp = await llm.ainvoke([
            SystemMessage(content=(
                "Based on the question and answer below, suggest exactly 3 short follow-up questions "
                "the user might want to ask next. Return ONLY the 3 questions, one per line, no numbering, "
                "no bullets, no extra text."
            )),
            HumanMessage(content=f"Question: {question}\n\nAnswer: {answer}"),
        ])
        lines = [l.strip() for l in resp.content.strip().split("\n") if l.strip()]
        return lines[:3]
    except Exception:
        logger.warning("Follow-up question generation failed", exc_info=True)
        return []


@app.get("/chats")
async def list_chats(request: Request, limit: int = 30):
    """List chat history entries for the current user."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    db = SessionLocal()
    chats = (
        db.query(ChatHistory)
        .filter(ChatHistory.user_id == user["user_id"])
        .order_by(ChatHistory.pinned.desc(), ChatHistory.created_at.desc())
        .limit(limit)
        .all()
    )
    db.close()
    return [
        {
            "id": c.id,
            "type": c.type,
            "title": c.title,
            "ref_id": c.ref_id,
            "pinned": bool(c.pinned),
            "created_at": c.created_at,
        }
        for c in chats
    ]


@app.patch("/chats/{chat_id}/pin")
async def toggle_pin_chat(chat_id: str, request: Request):
    """Toggle pin status of a chat."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    db = SessionLocal()
    try:
        chat = db.query(ChatHistory).filter(
            ChatHistory.id == chat_id,
            ChatHistory.user_id == user["user_id"],
        ).first()
        if not chat:
            raise HTTPException(status_code=404, detail="Chat not found")
        chat.pinned = not chat.pinned
        db.commit()
        return {"id": chat.id, "pinned": bool(chat.pinned)}
    except HTTPException:
        raise
    except Exception:
        logger.error("Failed to toggle pin for chat %s", chat_id, exc_info=True)
        db.rollback()
        raise HTTPException(status_code=500, detail="Failed to update pin status")
    finally:
        db.close()


@app.get("/chats/{chat_id}")
async def get_chat_detail(chat_id: str, request: Request):
    """Get full chat detail — resolves to Q&A or Research data."""
    user = get_current_user(request)
    if not user:
        raise HTTPException(status_code=401, detail="Not authenticated")
    db = SessionLocal()
    chat = db.query(ChatHistory).filter(ChatHistory.id == chat_id).first()
    if not chat or chat.user_id != user["user_id"]:
        db.close()
        raise HTTPException(status_code=404, detail="Chat not found")

    result = {"id": chat.id, "type": chat.type, "title": chat.title, "created_at": chat.created_at}

    if chat.type == "Q&A":
        # ref_id is now chat_session_id — fetch all messages in session
        qa_list = (
            db.query(QAInteraction)
            .filter(QAInteraction.chat_session_id == chat.ref_id)
            .order_by(QAInteraction.created_at.asc())
            .all()
        )
        if qa_list:
            result["chat_session_id"] = chat.ref_id
            result["messages"] = [
                {
                    "question": qa.question,
                    "answer": qa.answer,
                    "sources": json.loads(qa.sources) if qa.sources else [],
                    "document_id": qa.document_id,
                    "created_at": str(qa.created_at),
                }
                for qa in qa_list
            ]
        else:
            result["messages"] = []
    elif chat.type == "Research":
        job = db.query(ResearchJob).filter(ResearchJob.id == chat.ref_id).first()
        if job:
            result["query"] = job.query
            result["report"] = job.report
            result["status"] = job.status

    db.close()
    return result


# ── Export (Notion + Google Docs) ─────────────────────────────────────────────

@app.get("/export/notion/connect")
async def notion_connect(request: Request):
    """Redirect user to Notion OAuth consent screen."""
    user = require_auth(request)
    client_id = os.getenv("NOTION_CLIENT_ID")
    if not client_id:
        raise HTTPException(status_code=500, detail="Notion OAuth not configured")

    base_url = os.getenv("APP_BASE_URL", "http://localhost:3000")
    redirect_uri = f"{base_url}/api/export/notion/callback"
    state = user["user_id"]  # pass user_id in state for callback

    url = (
        f"https://api.notion.com/v1/oauth/authorize"
        f"?client_id={client_id}"
        f"&response_type=code"
        f"&owner=user"
        f"&redirect_uri={redirect_uri}"
        f"&state={state}"
    )
    return RedirectResponse(url=url)


@app.get("/export/notion/callback")
async def notion_callback(code: str, state: str = ""):
    """Handle Notion OAuth callback and store the token."""
    import httpx
    import base64

    client_id = os.getenv("NOTION_CLIENT_ID")
    client_secret = os.getenv("NOTION_CLIENT_SECRET")
    base_url = os.getenv("APP_BASE_URL", "http://localhost:3000")
    redirect_uri = f"{base_url}/api/export/notion/callback"

    # Exchange code for access token
    credentials = base64.b64encode(f"{client_id}:{client_secret}".encode()).decode()
    async with httpx.AsyncClient() as http:
        res = await http.post(
            "https://api.notion.com/v1/oauth/token",
            headers={
                "Authorization": f"Basic {credentials}",
                "Content-Type": "application/json",
            },
            json={
                "grant_type": "authorization_code",
                "code": code,
                "redirect_uri": redirect_uri,
            },
        )
        if res.status_code != 200:
            return RedirectResponse(url=f"{base_url}/?export_error=notion_auth_failed")
        data = res.json()

    access_token = data.get("access_token")
    workspace_id = data.get("workspace_id", "")
    user_id = state

    if not access_token or not user_id:
        return RedirectResponse(url=f"{base_url}/?export_error=notion_no_token")

    # Store or update connection
    db = SessionLocal()
    conn = db.query(OAuthConnection).filter(
        OAuthConnection.user_id == user_id,
        OAuthConnection.provider == "notion",
    ).first()
    if conn:
        conn.access_token = access_token
        conn.workspace_id = workspace_id
        conn.extra = json.dumps(data)
    else:
        conn = OAuthConnection(
            user_id=user_id,
            provider="notion",
            access_token=access_token,
            workspace_id=workspace_id,
            extra=json.dumps(data),
        )
        db.add(conn)
    db.commit()
    db.close()

    return RedirectResponse(url=f"{base_url}/?export_success=notion")


@app.get("/export/google-docs/connect")
async def google_docs_connect(request: Request):
    """Redirect to Google OAuth with Docs scope."""
    user = require_auth(request)
    client_id = os.getenv("GOOGLE_CLIENT_ID")
    if not client_id:
        raise HTTPException(status_code=500, detail="Google OAuth not configured")

    base_url = os.getenv("APP_BASE_URL", "http://localhost:3000")
    redirect_uri = f"{base_url}/api/export/google-docs/callback"
    scope = "https://www.googleapis.com/auth/documents https://www.googleapis.com/auth/drive.file"
    state = user["user_id"]

    url = (
        f"https://accounts.google.com/o/oauth2/v2/auth"
        f"?client_id={client_id}"
        f"&redirect_uri={redirect_uri}"
        f"&response_type=code"
        f"&scope={scope}"
        f"&access_type=offline"
        f"&prompt=consent"
        f"&state={state}"
    )
    return RedirectResponse(url=url)


@app.get("/export/google-docs/callback")
async def google_docs_callback(code: str, state: str = ""):
    """Handle Google Docs OAuth callback."""
    import httpx

    client_id = os.getenv("GOOGLE_CLIENT_ID")
    client_secret = os.getenv("GOOGLE_CLIENT_SECRET")
    base_url = os.getenv("APP_BASE_URL", "http://localhost:3000")
    redirect_uri = f"{base_url}/api/export/google-docs/callback"

    async with httpx.AsyncClient() as http:
        res = await http.post(
            "https://oauth2.googleapis.com/token",
            data={
                "code": code,
                "client_id": client_id,
                "client_secret": client_secret,
                "redirect_uri": redirect_uri,
                "grant_type": "authorization_code",
            },
        )
        if res.status_code != 200:
            return RedirectResponse(url=f"{base_url}/?export_error=google_docs_auth_failed")
        tokens = res.json()

    user_id = state
    if not user_id:
        return RedirectResponse(url=f"{base_url}/?export_error=google_docs_no_user")

    db = SessionLocal()
    conn = db.query(OAuthConnection).filter(
        OAuthConnection.user_id == user_id,
        OAuthConnection.provider == "google_docs",
    ).first()
    if conn:
        conn.access_token = tokens.get("access_token", "")
        conn.refresh_token = tokens.get("refresh_token", conn.refresh_token)
        conn.extra = json.dumps(tokens)
    else:
        conn = OAuthConnection(
            user_id=user_id,
            provider="google_docs",
            access_token=tokens.get("access_token", ""),
            refresh_token=tokens.get("refresh_token"),
            extra=json.dumps(tokens),
        )
        db.add(conn)
    db.commit()
    db.close()

    return RedirectResponse(url=f"{base_url}/?export_success=google_docs")


@app.get("/export/connections")
async def get_export_connections(request: Request):
    """Check which export services the user has connected."""
    user = require_auth(request)
    db = SessionLocal()
    connections = db.query(OAuthConnection).filter(
        OAuthConnection.user_id == user["user_id"]
    ).all()
    db.close()

    return {
        "notion": any(c.provider == "notion" for c in connections),
        "google_docs": any(c.provider == "google_docs" for c in connections),
    }


@app.post("/export/notion/{job_id}")
async def export_to_notion(job_id: str, request: Request):
    """Export a research report to Notion as a new page."""
    import httpx
    from shared.export_converters import markdown_to_notion_blocks

    user = require_auth(request)
    db = SessionLocal()

    job = db.query(ResearchJob).filter(ResearchJob.id == job_id).first()
    if not job or not job.report:
        db.close()
        raise HTTPException(status_code=404, detail="Report not found")

    conn = db.query(OAuthConnection).filter(
        OAuthConnection.user_id == user["user_id"],
        OAuthConnection.provider == "notion",
    ).first()
    db.close()

    if not conn:
        raise HTTPException(status_code=400, detail="Notion not connected. Please connect first.")

    blocks = markdown_to_notion_blocks(job.report)
    # Notion API limits to 100 blocks per request
    blocks = blocks[:100]

    title = f"Research: {job.query[:60]}"

    async with httpx.AsyncClient() as http:
        # Create a page in the user's workspace (as a top-level page)
        res = await http.post(
            "https://api.notion.com/v1/pages",
            headers={
                "Authorization": f"Bearer {conn.access_token}",
                "Content-Type": "application/json",
                "Notion-Version": "2022-06-28",
            },
            json={
                "parent": {"type": "workspace", "workspace": True},
                "properties": {
                    "title": [{"text": {"content": title}}],
                },
                "children": blocks,
            },
        )

        if res.status_code not in (200, 201):
            error_detail = res.json().get("message", res.text[:200])
            raise HTTPException(status_code=502, detail=f"Notion API error: {error_detail}")

        page_data = res.json()
        page_url = page_data.get("url", "")

    return {"success": True, "url": page_url, "title": title}


@app.post("/export/google-docs/{job_id}")
async def export_to_google_docs(job_id: str, request: Request):
    """Export a research report to a new Google Doc."""
    import httpx
    from shared.export_converters import markdown_to_google_docs_requests

    user = require_auth(request)
    db = SessionLocal()

    job = db.query(ResearchJob).filter(ResearchJob.id == job_id).first()
    if not job or not job.report:
        db.close()
        raise HTTPException(status_code=404, detail="Report not found")

    conn = db.query(OAuthConnection).filter(
        OAuthConnection.user_id == user["user_id"],
        OAuthConnection.provider == "google_docs",
    ).first()
    db.close()

    if not conn:
        raise HTTPException(status_code=400, detail="Google Docs not connected. Please connect first.")

    # Refresh token if we have a refresh_token
    access_token = conn.access_token
    if conn.refresh_token:
        async with httpx.AsyncClient() as http:
            refresh_res = await http.post(
                "https://oauth2.googleapis.com/token",
                data={
                    "client_id": os.getenv("GOOGLE_CLIENT_ID"),
                    "client_secret": os.getenv("GOOGLE_CLIENT_SECRET"),
                    "refresh_token": conn.refresh_token,
                    "grant_type": "refresh_token",
                },
            )
            if refresh_res.status_code == 200:
                new_tokens = refresh_res.json()
                access_token = new_tokens.get("access_token", access_token)
                # Update stored token
                db2 = SessionLocal()
                c = db2.query(OAuthConnection).filter(OAuthConnection.id == conn.id).first()
                if c:
                    c.access_token = access_token
                    db2.commit()
                db2.close()

    title = f"Research: {job.query[:60]}"

    async with httpx.AsyncClient() as http:
        # Create empty document
        create_res = await http.post(
            "https://docs.googleapis.com/v1/documents",
            headers={
                "Authorization": f"Bearer {access_token}",
                "Content-Type": "application/json",
            },
            json={"title": title},
        )

        if create_res.status_code not in (200, 201):
            logger.error(f"Google Docs create failed: {create_res.status_code} {create_res.text}")
            raise HTTPException(status_code=502, detail=f"Failed to create Google Doc: {create_res.json().get('error', {}).get('message', create_res.text[:200])}")

        doc = create_res.json()
        doc_id = doc["documentId"]

        # Insert content
        requests = markdown_to_google_docs_requests(job.report)
        if requests:
            batch_res = await http.post(
                f"https://docs.googleapis.com/v1/documents/{doc_id}:batchUpdate",
                headers={
                    "Authorization": f"Bearer {access_token}",
                    "Content-Type": "application/json",
                },
                json={"requests": requests},
            )
            if batch_res.status_code not in (200, 201):
                # Doc was created but content failed — still return the link
                pass

    doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"
    return {"success": True, "url": doc_url, "title": title}


# ── Notifications ─────────────────────────────────────────────────────────────

@app.get("/notifications")
async def list_notifications(request: Request, limit: int = 20):
    """List notification history for the current user."""
    user = require_auth(request)
    db = SessionLocal()
    notifs = (
        db.query(NotificationLog)
        .filter(NotificationLog.user_id == user["user_id"])
        .order_by(NotificationLog.created_at.desc())
        .limit(limit)
        .all()
    )
    db.close()
    unread = sum(1 for n in notifs if not n.is_read)
    return {
        "unread_count": unread,
        "notifications": [
            {
                "id": n.id,
                "type": n.type,
                "subject": n.subject,
                "preview": n.preview,
                "ref_id": n.ref_id,
                "status": n.status,
                "is_read": n.is_read,
                "created_at": n.created_at.isoformat() if n.created_at else None,
            }
            for n in notifs
        ],
    }


@app.patch("/notifications/{notif_id}/read")
async def mark_notification_read(notif_id: str, request: Request):
    """Mark a single notification as read."""
    user = require_auth(request)
    db = SessionLocal()
    try:
        notif = db.query(NotificationLog).filter(
            NotificationLog.id == notif_id,
            NotificationLog.user_id == user["user_id"],
        ).first()
        if not notif:
            raise HTTPException(status_code=404, detail="Notification not found")
        notif.is_read = True
        db.commit()
        return {"id": notif.id, "is_read": True}
    finally:
        db.close()


@app.post("/notifications/read-all")
async def mark_all_read(request: Request):
    """Mark all notifications as read for the current user."""
    user = require_auth(request)
    db = SessionLocal()
    try:
        db.query(NotificationLog).filter(
            NotificationLog.user_id == user["user_id"],
            NotificationLog.is_read == False,
        ).update({"is_read": True})
        db.commit()
        return {"success": True}
    finally:
        db.close()


# ── WebSocket — live agent status stream ─────────────────────────────────────

@app.websocket("/ws/{job_id}")
async def websocket_status(websocket: WebSocket, job_id: str):
    """
    Connect to receive real-time agent status updates for a job.
    Messages are JSON: {"job_id": "...", "agent": "search", "status": "running"}
    """
    await websocket.accept()
    channel = f"status:{job_id}"

    try:
        while True:
            try:
                message = await asyncio.wait_for(bus.consume(channel), timeout=30)
            except asyncio.TimeoutError:
                # Send ping to keep connection alive through nginx
                await websocket.send_text(json.dumps({"type": "ping"}))
                continue

            await websocket.send_text(message)

            # Close when orchestrator signals done or failed
            parsed = json.loads(message)
            if parsed.get("agent") == "orchestrator" and parsed.get("status") in ("done", "failed"):
                break

    except WebSocketDisconnect:
        logger.debug("WebSocket client disconnected for job %s", job_id)
    except Exception as e:
        logger.error("WebSocket error for job %s: %s", job_id, e, exc_info=True)
    finally:
        try:
            await websocket.close()
        except Exception:
            logger.debug("WebSocket already closed for job %s", job_id, exc_info=True)


@app.get("/health")
async def health():
    """Liveness probe used by Docker / load balancers."""
    return {"status": "ok"}


# ── Quick Q&A (RAG-powered, no full pipeline) ────────────────────────────────

class QuickAskRequest(BaseModel):
    question: str
    conversation_history: list = []  # [{"role": "user"|"assistant", "content": "..."}]
    chat_session_id: str | None = None


@app.post("/ask")
async def quick_ask(req: QuickAskRequest, request: Request):
    """
    Quick question answering using RAG context from past research.
    Supports multi-turn conversation via conversation_history.
    """
    if not req.question.strip():
        raise HTTPException(status_code=400, detail="Question cannot be empty")

    from shared.vector_store import retrieve_relevant_chunks
    from langchain_openai import AzureChatOpenAI
    from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
    import os

    user = get_current_user(request)

    # Retrieve relevant context from the knowledge base
    relevant = retrieve_relevant_chunks(query=req.question, n_results=8)

    context = ""
    sources = []
    if relevant:
        context = "\n\n".join(
            f"[{i+1}] {r['url']}\n{r['content']}"
            for i, r in enumerate(relevant)
        )
        sources = [
            {"url": r["url"], "distance": round(r["distance"], 3)}
            for r in relevant
        ]

    llm = AzureChatOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-12-01-preview"),
        max_tokens=1500,
    )

    system_prompt = (
        "You are a knowledgeable research assistant. Answer the user's question clearly and concisely. "
    )
    if context:
        system_prompt += (
            "Use the provided context from past research to inform your answer. "
            "Cite sources as [1], [2], etc. when applicable. "
            "If the context doesn't cover the question fully, supplement with your own knowledge but note it."
        )
    else:
        system_prompt += "No relevant past research was found, so answer from your general knowledge."

    # Build messages with conversation history
    messages = [SystemMessage(content=system_prompt)]
    for msg in req.conversation_history:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        elif msg["role"] == "assistant":
            messages.append(AIMessage(content=msg["content"]))

    user_content = req.question
    if context:
        user_content = f"Context from knowledge base:\n{context}\n\nQuestion: {req.question}"
    messages.append(HumanMessage(content=user_content))

    response = await llm.ainvoke(messages)
    answer_text = response.content.strip()

    # Generate follow-up questions
    followups = await _generate_followups(llm, req.question, answer_text)

    # Determine chat session
    chat_session_id = req.chat_session_id or str(uuid.uuid4())

    # Save Q&A interaction + chat history
    if user:
        db = SessionLocal()
        try:
            qa = QAInteraction(
                user_id=user["user_id"],
                chat_session_id=chat_session_id,
                document_id=None,
                question=req.question,
                answer=answer_text,
                sources=json.dumps(sources),
            )
            db.add(qa)
            db.commit()
            db.refresh(qa)
            # Only create chat history entry for first message in session
            if not req.chat_session_id:
                title = await _generate_chat_title(req.question)
                _save_chat_history(user["user_id"], "Q&A", title, chat_session_id)
        except Exception:
            logger.error("Failed to persist Q&A interaction for session %s", chat_session_id, exc_info=True)
            db.rollback()
        finally:
            db.close()

    return {
        "answer": answer_text,
        "sources": sources,
        "has_context": bool(relevant),
        "followups": followups,
        "chat_session_id": chat_session_id,
    }


@app.post("/ask/stream")
async def quick_ask_stream(req: QuickAskRequest, request: Request):
    """Stream Q&A answer token-by-token via SSE."""
    if not req.question.strip():
        raise HTTPException(status_code=400, detail="Question cannot be empty")

    from shared.vector_store import retrieve_relevant_chunks
    from langchain_openai import AzureChatOpenAI
    from langchain_core.messages import SystemMessage, HumanMessage, AIMessage
    import os

    user = get_current_user(request)

    relevant = retrieve_relevant_chunks(query=req.question, n_results=8)
    context = ""
    sources = []
    if relevant:
        context = "\n\n".join(
            f"[{i+1}] {r['url']}\n{r['content']}"
            for i, r in enumerate(relevant)
        )
        sources = [
            {"url": r["url"], "distance": round(r["distance"], 3)}
            for r in relevant
        ]

    llm = AzureChatOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-12-01-preview"),
        max_tokens=1500,
        streaming=True,
    )

    system_prompt = (
        "You are a knowledgeable research assistant. Answer the user's question clearly and concisely. "
    )
    if context:
        system_prompt += (
            "Use the provided context from past research to inform your answer. "
            "Cite sources as [1], [2], etc. when applicable. "
            "If the context doesn't cover the question fully, supplement with your own knowledge but note it."
        )
    else:
        system_prompt += "No relevant past research was found, so answer from your general knowledge."

    messages = [SystemMessage(content=system_prompt)]
    for msg in req.conversation_history:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        elif msg["role"] == "assistant":
            messages.append(AIMessage(content=msg["content"]))

    user_content = req.question
    if context:
        user_content = f"Context from knowledge base:\n{context}\n\nQuestion: {req.question}"
    messages.append(HumanMessage(content=user_content))

    chat_session_id = req.chat_session_id or str(uuid.uuid4())

    async def event_generator():
        answer_text = ""
        try:
            # Send metadata first
            yield f"data: {json.dumps({'type': 'meta', 'sources': sources, 'has_context': bool(relevant), 'chat_session_id': chat_session_id})}\n\n"

            async for chunk in llm.astream(messages):
                token = chunk.content
                if token:
                    answer_text += token
                    yield f"data: {json.dumps({'type': 'token', 'token': token})}\n\n"

            # Generate follow-ups after streaming completes
            followups = await _generate_followups(llm, req.question, answer_text)
            yield f"data: {json.dumps({'type': 'followups', 'followups': followups})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            logger.error("Quick-ask stream failed: %s", e, exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"

        # Save to DB after stream completes
        if user and answer_text:
            db = SessionLocal()
            try:
                qa = QAInteraction(
                    user_id=user["user_id"],
                    chat_session_id=chat_session_id,
                    document_id=None,
                    question=req.question,
                    answer=answer_text,
                    sources=json.dumps(sources),
                )
                db.add(qa)
                db.commit()
                if not req.chat_session_id:
                    title = await _generate_chat_title(req.question)
                    _save_chat_history(user["user_id"], "Q&A", title, chat_session_id)
            except Exception:
                logger.error("Failed to persist streamed quick-ask Q&A for session %s", chat_session_id, exc_info=True)
                db.rollback()
            finally:
                db.close()

    return StreamingResponse(event_generator(), media_type="text/event-stream")


# ── Document Upload + Q&A (RAG) ──────────────────────────────────────────────

from shared.document_processor import extract_text, chunk_text, SUPPORTED_EXTENSIONS
from shared.vector_store import store_document_chunks, query_document, embed_texts
from langchain_openai import AzureChatOpenAI
from langchain_core.messages import SystemMessage, HumanMessage
import os

_llm = None

def _get_llm():
    """Lazily build and cache a shared AzureChatOpenAI client for document Q&A."""
    global _llm
    if _llm is None:
        _llm = AzureChatOpenAI(
            azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
            openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
            azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
            openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-12-01-preview"),
            max_tokens=1500,
        )
    return _llm


# In-memory doc registry (maps doc_id -> filename)
_documents: dict[str, dict] = {}


@app.post("/documents/upload")
async def upload_document(request: Request, file: UploadFile = File(...)):
    """Upload a document for RAG-based Q&A."""
    from pathlib import Path

    user = get_current_user(request)
    user_id = user["user_id"] if user else None

    ext = Path(file.filename).suffix.lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {ext}. Supported: {', '.join(SUPPORTED_EXTENSIONS)}",
        )

    content = await file.read()
    if len(content) > 20 * 1024 * 1024:  # 20MB limit
        raise HTTPException(status_code=400, detail="File too large (max 20MB)")

    doc_id = str(uuid.uuid4())

    try:
        text = extract_text(file.filename, content)
        chunks = chunk_text(text)

        if not chunks:
            raise HTTPException(status_code=400, detail="Could not extract text from file")

        stored = store_document_chunks(chunks, doc_id, file.filename)

        # Persist to DB
        db = SessionLocal()
        doc = Document(
            id=doc_id,
            user_id=user_id,
            filename=file.filename,
            chunks_count=stored,
            text_length=len(text),
        )
        db.add(doc)
        db.commit()
        db.close()

        return {
            "doc_id": doc_id,
            "filename": file.filename,
            "chunks_stored": stored,
            "text_length": len(text),
        }

    except ValueError as e:
        logger.warning("Document upload validation failed: %s", e)
        raise HTTPException(status_code=400, detail=str(e))


class QuestionRequest(BaseModel):
    question: str
    doc_id: str
    conversation_history: list = []
    chat_session_id: str | None = None


@app.post("/documents/ask")
async def ask_document(req: QuestionRequest, request: Request):
    """Ask a question about an uploaded document using RAG."""
    if not req.question.strip():
        raise HTTPException(status_code=400, detail="Question cannot be empty")

    user = get_current_user(request)

    # Retrieve relevant chunks
    relevant_chunks = query_document(
        query=req.question,
        doc_id=req.doc_id,
        n_results=6,
    )

    if not relevant_chunks:
        raise HTTPException(
            status_code=404,
            detail="No document found with this ID or no relevant content",
        )

    # Build context from retrieved chunks
    context = "\n\n".join(
        f"[Chunk {r['chunk_index'] + 1}]: {r['content']}"
        for r in relevant_chunks
    )

    from langchain_core.messages import AIMessage

    # Ask GPT with the retrieved context + conversation history
    llm = _get_llm()
    messages = [
        SystemMessage(content=(
            "You are a helpful assistant that answers questions based on the provided document context. "
            "Only answer based on the context given. If the answer is not in the context, say so. "
            "Be concise and accurate. Do not mention chunk numbers or internal references."
        )),
    ]
    for msg in req.conversation_history:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        elif msg["role"] == "assistant":
            messages.append(AIMessage(content=msg["content"]))
    messages.append(HumanMessage(content=f"Context from document:\n{context}\n\nQuestion: {req.question}"))

    response = await llm.ainvoke(messages)

    answer_text = response.content.strip()
    sources_list = [
        {"chunk_index": r["chunk_index"], "distance": round(r["distance"], 3)}
        for r in relevant_chunks
    ]

    # Generate follow-up questions
    followups = await _generate_followups(llm, req.question, answer_text)

    # Determine chat session
    chat_session_id = req.chat_session_id or str(uuid.uuid4())

    # Save Q&A interaction + chat history
    if user:
        db = SessionLocal()
        try:
            qa = QAInteraction(
                user_id=user["user_id"],
                chat_session_id=chat_session_id,
                document_id=req.doc_id,
                question=req.question,
                answer=answer_text,
                sources=json.dumps(sources_list),
            )
            db.add(qa)
            db.commit()
            db.refresh(qa)
            if not req.chat_session_id:
                title = await _generate_chat_title(req.question)
                _save_chat_history(user["user_id"], "Q&A", title, chat_session_id)
        except Exception:
            logger.error("Failed to persist document Q&A for session %s", chat_session_id, exc_info=True)
            db.rollback()
        finally:
            db.close()

    return {
        "answer": answer_text,
        "sources": sources_list,
        "followups": followups,
        "chat_session_id": chat_session_id,
    }


@app.post("/documents/ask/stream")
async def ask_document_stream(req: QuestionRequest, request: Request):
    """Stream document Q&A answer token-by-token via SSE. Falls back to web search if not in document."""
    if not req.question.strip():
        raise HTTPException(status_code=400, detail="Question cannot be empty")

    user = get_current_user(request)

    relevant_chunks = query_document(
        query=req.question,
        doc_id=req.doc_id,
        n_results=6,
    )

    if not relevant_chunks:
        raise HTTPException(
            status_code=404,
            detail="No document found with this ID or no relevant content",
        )

    # Determine if chunks are relevant enough (low distance = high relevance)
    avg_distance = sum(r["distance"] for r in relevant_chunks) / len(relevant_chunks)
    doc_is_relevant = avg_distance < 1.2  # threshold for "found in document"

    web_context = ""
    web_sources = []
    if not doc_is_relevant:
        # Fall back to web search
        try:
            from tavily import TavilyClient
            tavily_client = TavilyClient(api_key=os.getenv("TAVILY_API_KEY"))
            web_results = tavily_client.search(query=req.question, max_results=5)
            web_items = web_results.get("results", [])
            if web_items:
                web_context = "\n\n".join(
                    f"[{i+1}] {r.get('url', '')}\n{r.get('content', '')}"
                    for i, r in enumerate(web_items)
                )
                web_sources = [{"url": r.get("url", ""), "title": r.get("title", "")} for r in web_items]
        except Exception as e:
            logger.warning("[doc_ask] Web search fallback failed: %s", e, exc_info=True)

    context = "\n\n".join(
        f"[Chunk {r['chunk_index'] + 1}]: {r['content']}"
        for r in relevant_chunks
    )

    sources_list = [
        {"chunk_index": r["chunk_index"], "distance": round(r["distance"], 3)}
        for r in relevant_chunks
    ]

    from langchain_core.messages import AIMessage as _AIMsg

    llm = AzureChatOpenAI(
        azure_endpoint=os.getenv("AZURE_OPENAI_ENDPOINT"),
        openai_api_key=os.getenv("AZURE_OPENAI_API_KEY"),
        azure_deployment=os.getenv("AZURE_OPENAI_DEPLOYMENT"),
        openai_api_version=os.getenv("AZURE_OPENAI_API_VERSION", "2024-12-01-preview"),
        max_tokens=1500,
        streaming=True,
    )

    if doc_is_relevant:
        system_prompt = (
            "You are a helpful assistant that answers questions based on the provided document context. "
            "Only answer based on the context given. If the answer is not in the context, say so. "
            "Be concise and accurate. Do not mention chunk numbers or internal references."
        )
        user_content = f"Context from document:\n{context}\n\nQuestion: {req.question}"
    else:
        system_prompt = (
            "You are a helpful research assistant. The user's question was not found in their attached document, "
            "so you are answering from web search results instead. "
            "Clearly indicate that the answer comes from web sources, not the document. "
            "Cite sources as [1], [2], etc. Be concise and accurate."
        )
        user_content = f"Web search results:\n{web_context}\n\nQuestion: {req.question}" if web_context else f"Question: {req.question}"

    messages = [SystemMessage(content=system_prompt)]
    for msg in req.conversation_history:
        if msg["role"] == "user":
            messages.append(HumanMessage(content=msg["content"]))
        elif msg["role"] == "assistant":
            messages.append(_AIMsg(content=msg["content"]))
    messages.append(HumanMessage(content=user_content))

    chat_session_id = req.chat_session_id or str(uuid.uuid4())

    async def event_generator():
        answer_text = ""
        try:
            # If falling back to web search, notify the UI
            if not doc_is_relevant:
                yield f"data: {json.dumps({'type': 'stage', 'stage': 'web_search', 'message': 'Not found in document — searching the web...'})}\n\n"

            final_sources = web_sources if not doc_is_relevant else sources_list
            yield f"data: {json.dumps({'type': 'meta', 'sources': final_sources, 'chat_session_id': chat_session_id, 'source_type': 'web' if not doc_is_relevant else 'document'})}\n\n"

            async for chunk in llm.astream(messages):
                token = chunk.content
                if token:
                    answer_text += token
                    yield f"data: {json.dumps({'type': 'token', 'token': token})}\n\n"

            followups = await _generate_followups(llm, req.question, answer_text)
            yield f"data: {json.dumps({'type': 'followups', 'followups': followups})}\n\n"
            yield f"data: {json.dumps({'type': 'done'})}\n\n"
        except Exception as e:
            logger.error("Document-ask stream failed for doc %s: %s", req.doc_id, e, exc_info=True)
            yield f"data: {json.dumps({'type': 'error', 'error': str(e)})}\n\n"

        if user and answer_text:
            db = SessionLocal()
            try:
                qa = QAInteraction(
                    user_id=user["user_id"],
                    chat_session_id=chat_session_id,
                    document_id=req.doc_id,
                    question=req.question,
                    answer=answer_text,
                    sources=json.dumps(final_sources),
                )
                db.add(qa)
                db.commit()
                if not req.chat_session_id:
                    title = await _generate_chat_title(req.question)
                    _save_chat_history(user["user_id"], "Q&A", title, chat_session_id)
            except Exception:
                logger.error("Failed to persist streamed document Q&A for session %s", chat_session_id, exc_info=True)
                db.rollback()
            finally:
                db.close()

    return StreamingResponse(event_generator(), media_type="text/event-stream")


@app.get("/documents")
async def list_documents(request: Request):
    """List uploaded documents for the current user."""
    user = get_current_user(request)
    db = SessionLocal()
    query = db.query(Document).order_by(Document.created_at.desc())
    if user:
        query = query.filter(Document.user_id == user["user_id"])
    docs = query.limit(20).all()
    db.close()
    return [
        {
            "doc_id": d.id,
            "filename": d.filename,
            "chunks_count": int(d.chunks_count),
            "text_length": int(d.text_length),
            "created_at": d.created_at,
        }
        for d in docs
    ]
